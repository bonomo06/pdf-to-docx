"""
Módulo de conversão de PDF para DOCX.
Implementa a lógica de conversão com validações de segurança.
"""

import os
import uuid
import tempfile
import logging
from pathlib import Path
from typing import BinaryIO

import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class PDFConverterError(Exception):
    """Exceção customizada para erros de conversão."""
    pass


class PDFToDocxConverter:
    """
    Classe para conversão de PDF para DOCX usando pdfplumber + python-docx.
    Focado em extração de tabelas com bordas visíveis.
    """
    
    # Magic bytes para validação de PDF
    PDF_MAGIC_BYTES = b"%PDF"
    
    def __init__(self, temp_dir: str | None = None):
        """
        Inicializa o conversor.
        
        Args:
            temp_dir: Diretório para arquivos temporários (opcional)
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        
    def _generate_safe_filename(self, extension: str) -> str:
        """
        Gera um nome de arquivo seguro e único.
        
        Args:
            extension: Extensão do arquivo (sem ponto)
            
        Returns:
            Nome de arquivo único
        """
        return f"{uuid.uuid4().hex}.{extension}"
    
    def _validate_pdf_content(self, content: bytes) -> bool:
        """
        Valida se o conteúdo é realmente um PDF.
        
        Args:
            content: Bytes do arquivo
            
        Returns:
            True se for um PDF válido
        """
        if len(content) < 4:
            return False
        return content[:4] == self.PDF_MAGIC_BYTES
    
    async def convert_from_bytes(self, pdf_content: bytes) -> bytes:
        """
        Converte PDF de bytes para DOCX.
        
        Args:
            pdf_content: Conteúdo do PDF em bytes
            
        Returns:
            Conteúdo do DOCX em bytes
        """
        # Valida o conteúdo do PDF
        if not self._validate_pdf_content(pdf_content):
            raise PDFConverterError("Arquivo não é um PDF válido")
        
        # Cria arquivos temporários
        pdf_filename = self._generate_safe_filename("pdf")
        docx_filename = self._generate_safe_filename("docx")
        
        pdf_path = Path(self.temp_dir) / pdf_filename
        docx_path = Path(self.temp_dir) / docx_filename
        
        try:
            # Escreve o PDF temporário
            pdf_path.write_bytes(pdf_content)
            
            logger.info(f"Iniciando conversão (pdfplumber): {pdf_filename} -> {docx_filename}")
            
            # --- Início da lógica pdfplumber ---
            doc = Document()
            
            # Tenta abrir o PDF e extrair tabelas
            with pdfplumber.open(str(pdf_path)) as pdf:
                for page in pdf.pages:
                    # Extrai tabelas da página
                    # table_settings pode ser ajustado se necessário (vertical_strategy, horizontal_strategy)
                    tables = page.extract_tables()
                    
                    if not tables:
                        # Fallback: se não achar tabela, tenta pegar o texto puro
                        text = page.extract_text()
                        if text:
                            doc.add_paragraph(text)
                        continue

                    for table_data in tables:
                        if not table_data:
                            continue
                            
                        rows = len(table_data)
                        # Garante número de colunas consistente
                        cols = max(len(row) for row in table_data) if rows > 0 else 0
                        
                        if rows == 0 or cols == 0:
                            continue

                        # Cria tabela no docx
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'  # FORÇA bordas visíveis
                        
                        for i, row in enumerate(table_data):
                            for j, cell_text in enumerate(row):
                                # Previne index error se a linha for menor que o max de colunas
                                if j < len(table.rows[i].cells):
                                    cell = table.rows[i].cells[j]
                                    cell.text = cell_text or ""
                    
                    # Adiciona quebra de página entre páginas do PDF, se não for a última
                    if page.page_number < len(pdf.pages):
                        doc.add_page_break()

            # Salva o arquivo DOCX
            doc.save(str(docx_path))
            # --- Fim da lógica pdfplumber ---
            
            # Lê o DOCX resultante
            if not docx_path.exists():
                raise PDFConverterError("Falha na conversão: arquivo DOCX não foi gerado")
            
            docx_content = docx_path.read_bytes()
            
            logger.info(f"Conversão concluída com sucesso: {len(docx_content)} bytes")
            
            return docx_content
            
        except Exception as e:
            logger.error(f"Erro na conversão: {str(e)}")
            raise PDFConverterError(f"Erro durante a conversão: {str(e)}")
        finally:
            # Limpa arquivos temporários
            self._safe_delete(pdf_path)
            self._safe_delete(docx_path)
    
    def _safe_delete(self, path: Path) -> None:
        """
        Remove arquivo de forma segura.
        """
        try:
            if path.exists():
                path.unlink()
        except Exception as e:
            logger.warning(f"Não foi possível remover arquivo temporário {path}: {e}")

# Instância global do conversor
converter = PDFToDocxConverter()
